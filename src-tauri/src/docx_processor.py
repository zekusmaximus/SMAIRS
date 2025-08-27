#!/usr/bin/env python3
"""
DOCX Track Changes Processor
Uses python-docx to add proper track changes to Word documents
"""

import sys
import json
import os
from datetime import datetime
from typing import Dict, List, Any, Optional

try:
    from docx import Document
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import OxmlElement, parse_xml
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install with: pip install python-docx", file=sys.stderr)

def create_track_changes_document(original_text: str, changes: List[Dict[str, Any]], metadata: Dict[str, Any]) -> str:
    """
    Create a DOCX document with proper track changes
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not available")
    
    # Create new document
    doc = Document()
    
    # Enable track changes in document settings
    enable_track_changes(doc)
    
    # Set document metadata
    set_document_metadata(doc, metadata)
    
    # Process text with track changes
    processed_text = apply_track_changes_to_text(original_text, changes)
    
    # Add content to document
    add_tracked_content(doc, processed_text, changes)
    
    # Save to temporary file
    output_path = os.path.join(os.path.dirname(__file__), '..', '..', 'out', 'track_changes.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    return os.path.abspath(output_path)

def enable_track_changes(doc: 'Document') -> None:
    """Enable track changes in document settings"""
    settings = doc.settings
    
    # Add track changes setting
    track_changes_xml = f'''
    <w:settings {nsdecls('w')}>
        <w:trackRevisions w:val="1"/>
        <w:doNotPromptForConvert w:val="1"/>
        <w:defaultTabStop w:val="708"/>
    </w:settings>
    '''
    
    # This is a simplified approach - in practice you'd need to properly merge with existing settings
    try:
        settings_elem = parse_xml(track_changes_xml)
        settings.element.getparent().replace(settings.element, settings_elem)
    except Exception as e:
        print(f"Warning: Could not set track changes in document settings: {e}", file=sys.stderr)

def set_document_metadata(doc: 'Document', metadata: Dict[str, Any]) -> None:
    """Set document core properties"""
    core_props = doc.core_properties
    
    if 'title' in metadata:
        core_props.title = metadata['title']
    if 'author' in metadata:
        if isinstance(metadata['author'], list):
            core_props.author = ', '.join(metadata['author'])
        else:
            core_props.author = metadata['author']
    if 'subject' in metadata:
        core_props.subject = metadata['subject']
    
    core_props.created = datetime.now()
    core_props.modified = datetime.now()

def apply_track_changes_to_text(original_text: str, changes: List[Dict[str, Any]]) -> str:
    """Apply track changes markers to text"""
    # Sort changes by position (descending to avoid offset issues)
    sorted_changes = sorted(changes, key=lambda x: x['position'], reverse=True)
    
    result = original_text
    change_markers = {}
    
    for change in sorted_changes:
        pos = change['position']
        change_type = change['type']
        content = change['content']
        change_id = change['id']
        
        # Store change metadata for later OOXML generation
        change_markers[change_id] = {
            'type': change_type,
            'author': change.get('author', 'SMAIRS'),
            'date': change.get('date', datetime.now().isoformat()),
            'comment': change.get('comment', '')
        }
        
        if change_type == 'insertion':
            # Insert with marker
            marker = f"[TRACK_INS_{change_id}]{content}[/TRACK_INS_{change_id}]"
            result = result[:pos] + marker + result[pos:]
        elif change_type == 'deletion':
            # Replace with deletion marker
            end_pos = pos + len(content)
            marker = f"[TRACK_DEL_{change_id}]{content}[/TRACK_DEL_{change_id}]"
            result = result[:pos] + marker + result[end_pos:]
    
    return result

def add_tracked_content(doc: 'Document', text: str, changes: List[Dict[str, Any]]) -> None:
    """Add content with track changes to the document"""
    
    # Split text into paragraphs
    paragraphs = text.split('\n\n')
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
            
        para = doc.add_paragraph()
        add_tracked_runs_to_paragraph(para, para_text, changes)

def add_tracked_runs_to_paragraph(para, text: str, changes: List[Dict[str, Any]]) -> None:
    """Add runs with track changes to a paragraph"""
    import re
    
    # Find all track change markers
    pattern = r'\[TRACK_(INS|DEL)_([^\]]+)\](.*?)\[/TRACK_\1_\2\]'
    
    last_end = 0
    
    for match in re.finditer(pattern, text):
        # Add normal text before this change
        if match.start() > last_end:
            normal_text = text[last_end:match.start()]
            if normal_text:
                para.add_run(normal_text)
        
        change_type = match.group(1).lower()
        change_id = match.group(2)
        content = match.group(3)
        
        # Find the corresponding change metadata
        change_meta = next((c for c in changes if c['id'] == change_id), None)
        if not change_meta:
            change_meta = {
                'author': 'SMAIRS',
                'date': datetime.now().isoformat(),
                'comment': ''
            }
        
        # Add tracked run
        if change_type == 'ins':
            add_insertion_run(para, content, change_id, change_meta)
        elif change_type == 'del':
            add_deletion_run(para, content, change_id, change_meta)
        
        last_end = match.end()
    
    # Add remaining normal text
    if last_end < len(text):
        remaining_text = text[last_end:]
        if remaining_text:
            para.add_run(remaining_text)

def add_insertion_run(para, content: str, change_id: str, change_meta: Dict[str, Any]) -> None:
    """Add an insertion track change run"""
    run = para.add_run(content)
    
    # Create insertion OOXML
    ins_xml = f'''
    <w:ins w:id="{change_id}" w:author="{escape_xml(change_meta['author'])}" w:date="{change_meta['date']}">
        <w:r>
            <w:t>{escape_xml(content)}</w:t>
        </w:r>
    </w:ins>
    '''
    
    try:
        # This is a simplified approach - proper implementation would need more complex OOXML manipulation
        ins_elem = parse_xml(ins_xml)
        run._element.getparent().replace(run._element, ins_elem)
    except Exception as e:
        print(f"Warning: Could not add insertion tracking to run: {e}", file=sys.stderr)
        # Fallback: just highlight the text
        run.font.highlight_color = 3  # Bright green

def add_deletion_run(para, content: str, change_id: str, change_meta: Dict[str, Any]) -> None:
    """Add a deletion track change run"""
    run = para.add_run(content)
    
    # Create deletion OOXML
    del_xml = f'''
    <w:del w:id="{change_id}" w:author="{escape_xml(change_meta['author'])}" w:date="{change_meta['date']}">
        <w:r>
            <w:delText>{escape_xml(content)}</w:delText>
        </w:r>
    </w:del>
    '''
    
    try:
        del_elem = parse_xml(del_xml)
        run._element.getparent().replace(run._element, del_elem)
    except Exception as e:
        print(f"Warning: Could not add deletion tracking to run: {e}", file=sys.stderr)
        # Fallback: strike through and red highlight
        run.font.strike = True
        run.font.highlight_color = 7  # Light red

def escape_xml(text: str) -> str:
    """Escape XML special characters"""
    return (text.replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&apos;'))

def main():
    """Command line interface"""
    if len(sys.argv) < 2:
        print("Usage: python docx_processor.py <json_input_file>", file=sys.stderr)
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        original_text = data.get('originalText', '')
        changes = data.get('changes', [])
        metadata = data.get('metadata', {})
        
        output_path = create_track_changes_document(original_text, changes, metadata)
        print(output_path)
        
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    main()